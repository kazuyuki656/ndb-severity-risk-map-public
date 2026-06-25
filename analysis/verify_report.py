import docx
from pathlib import Path

output_dir = Path(__file__).resolve().parent / 'output'
p = output_dir / 'NDB重症化予防ニーズマップ_分析レポート.docx'
d = docx.Document(p)
out = output_dir / 'report_verify.txt'
with open(out, 'w', encoding='utf-8') as f:
    f.write(f"段落数: {len(d.paragraphs)}\n")
    f.write(f"表の数: {len(d.tables)}\n")
    n_img = sum(1 for r in d.part.rels.values() if 'image' in r.reltype)
    f.write(f"画像数: {n_img}\n\n")
    f.write("=== 見出し一覧 ===\n")
    for para in d.paragraphs:
        sn = para.style.name if para.style else ''
        if sn and sn.startswith('Heading'):
            f.write(f"[{sn}] {para.text}\n")
    f.write("\n=== 各表のサイズ ===\n")
    for i, t in enumerate(d.tables):
        hdr = [c.text for c in t.rows[0].cells]
        f.write(f"表{i+1}: {len(t.rows)}行 x {len(t.columns)}列  ヘッダ: {hdr}\n")
